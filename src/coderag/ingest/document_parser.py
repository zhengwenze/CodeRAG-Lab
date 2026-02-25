from typing import List, Dict, Any, Optional
from pathlib import Path
import logging

logger = logging.getLogger(__name__)


class DocumentParser:
    """文档解析器，支持多种格式"""

    PDF = "pdf"
    MARKDOWN = "md"
    WORD = "docx"
    WORD_OLD = "doc"    
    TEXT = "txt"
    CSV = "csv"
    JSON = "json"
    YAML = "yaml"
    YML = "yml"
    HTML = "html"
    XML = "xml"

    SUPPORTED_FORMATS = {PDF, MARKDOWN, WORD, WORD_OLD, TEXT, CSV, JSON, YAML, YML, HTML, XML}

    @staticmethod
    def is_supported(file_path: str) -> bool:
        """检查文件格式是否支持"""
        ext = Path(file_path).suffix.lower().lstrip('.')
        return ext in DocumentParser.SUPPORTED_FORMATS

    @staticmethod
    def get_extension(file_path: str) -> str:
        """获取文件扩展名"""
        return Path(file_path).suffix.lower().lstrip('.')

    @classmethod
    def parse(cls, file_path: str, content: Optional[str] = None) -> str:
        """
        解析文档，返回纯文本内容
        
        Args:
            file_path: 文件路径
            content: 如果已读取内容则传入，否则从文件读取
        
        Returns:
            解析后的纯文本内容
        """
        ext = cls.get_extension(file_path)
        
        if ext == cls.PDF:
            return cls._parse_pdf(file_path, content)
        elif ext in {cls.MARKDOWN, cls.TEXT, cls.CSV, cls.JSON, cls.YAML, cls.YML, cls.HTML, cls.XML}:
            return cls._parse_text_based(file_path, content)
        elif ext in {cls.WORD, cls.WORD_OLD}:
            return cls._parse_word(file_path, content)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")

    @classmethod
    def _parse_pdf(cls, file_path: str, content: Optional[str] = None) -> str:
        """解析 PDF 文件"""
        try:
            from pypdf import PdfReader
        except ImportError:
            logger.warning("pypdf 未安装，将使用 fallback 方式解析")
            return cls._parse_text_based(file_path, content)

        if content is not None:
            raise ValueError("PDF 解析需要文件路径，不支持直接传入内容")

        try:
            reader = PdfReader(file_path)
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"PDF 解析失败: {file_path}, error: {e}")
            raise

    @classmethod
    def _parse_word(cls, file_path: str, content: Optional[str] = None) -> str:
        """解析 Word 文件 (.docx, .doc)"""
        try:
            from docx import Document
        except ImportError:
            logger.warning("python-docx 未安装，将使用 fallback 方式解析")
            return cls._parse_text_based(file_path, content)

        if content is not None:
            raise ValueError("Word 解析需要文件路径，不支持直接传入内容")

        ext = cls.get_extension(file_path)
        
        try:
            doc = Document(file_path)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Word 解析失败: {file_path}, error: {e}")
            raise

    @classmethod
    def _parse_text_based(cls, file_path: str, content: Optional[str] = None) -> str:
        """
        解析基于文本的格式 (txt, md, csv, json, yaml, html, xml)
        """
        if content is not None:
            return content
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='gbk') as f:
                    return f.read()
            except Exception as e:
                logger.error(f"文件读取失败: {file_path}, error: {e}")
                raise

    @classmethod
    def get_metadata(cls, file_path: str) -> Dict[str, Any]:
        """
        获取文档元信息
        
        Returns:
            包含元信息的字典
        """
        path = Path(file_path)
        ext = cls.get_extension(file_path)
        
        metadata = {
            "file_name": path.name,
            "file_extension": ext,
            "file_size": path.stat().st_size if path.exists() else 0,
            "file_type": cls._get_file_type(ext),
        }
        
        return metadata

    @staticmethod
    def _get_file_type(ext: str) -> str:
        """根据扩展名获取文件类型"""
        type_mapping = {
            DocumentParser.PDF: "PDF文档",
            DocumentParser.MARKDOWN: "Markdown文档",
            DocumentParser.WORD: "Word文档",
            DocumentParser.WORD_OLD: "Word文档",
            DocumentParser.TEXT: "文本文件",
            DocumentParser.CSV: "CSV表格",
            DocumentParser.JSON: "JSON数据",
            DocumentParser.YAML: "YAML配置",
            DocumentParser.YML: "YAML配置",
            DocumentParser.HTML: "HTML网页",
            DocumentParser.XML: "XML数据",
        }
        return type_mapping.get(ext, "未知类型")
